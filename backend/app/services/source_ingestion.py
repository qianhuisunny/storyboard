"""Bounded, SSRF-safe primitives for project source ingestion."""

from __future__ import annotations

import asyncio
import ipaddress
import multiprocessing
import os
import queue
import socket
import ssl
import sys
import threading
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Awaitable, Callable, Literal, Optional
from urllib.parse import urljoin, urlsplit

import httpx
import httpcore
from httpcore._backends.auto import AutoBackend


MAX_FETCH_BYTES = 2 * 1024 * 1024
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARS = 50_000
MAX_EXTRACTED_BYTES = 200_000
MAX_DOCX_FILES = 1_000
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_PDF_PAGES = 200
MAX_PDF_OBJECTS = 50_000
MAX_DOCX_PARAGRAPHS = 20_000
MAX_REDIRECTS = 4
SAFE_CONTENT_TYPES = {"text/html", "text/plain", "application/xhtml+xml"}
REDIRECT_STATUSES = {301, 302, 303, 307, 308}

Resolver = Callable[[str, int], Awaitable[list[str]]]


class SourceIngestionError(ValueError):
    """Raised when a source is unsafe, unsupported, or exceeds a bound."""


@dataclass(frozen=True)
class FetchedText:
    final_url: str
    content_type: str
    text: str


@dataclass(frozen=True)
class ValidatedTarget:
    url: str
    hostname: str
    pinned_ip: str


class PinnedNetworkBackend:
    """Replace only TCP DNS resolution; httpcore retains the origin and TLS SNI."""

    def __init__(
        self,
        *,
        original_hostname: str,
        pinned_ip: str,
        delegate=None,
    ) -> None:
        self.original_hostname = original_hostname.rstrip(".").lower()
        self.pinned_ip = pinned_ip
        self.delegate = delegate or AutoBackend()

    async def connect_tcp(
        self,
        host: str,
        port: int,
        timeout: float | None = None,
        local_address: str | None = None,
        socket_options=None,
    ):
        normalized = host.rstrip(".").lower()
        if normalized != self.original_hostname:
            raise SourceIngestionError("Network target did not match the validated host")
        return await self.delegate.connect_tcp(
            self.pinned_ip,
            port,
            timeout=timeout,
            local_address=local_address,
            socket_options=socket_options,
        )

    async def connect_unix_socket(self, path: str, **kwargs):
        return await self.delegate.connect_unix_socket(path, **kwargs)

    async def sleep(self, seconds: float) -> None:
        await self.delegate.sleep(seconds)


class PinnedAsyncHTTPTransport(httpx.AsyncHTTPTransport):
    """httpx transport whose socket connects to a single validated address."""

    def __init__(self, *, hostname: str, pinned_ip: str) -> None:
        super().__init__(trust_env=False)
        self._pool = httpcore.AsyncConnectionPool(
            ssl_context=ssl.create_default_context(),
            max_connections=1,
            max_keepalive_connections=0,
            retries=0,
            network_backend=PinnedNetworkBackend(
                original_hostname=hostname,
                pinned_ip=pinned_ip,
            ),
        )


def truncate_utf8(value: str, max_bytes: int = MAX_EXTRACTED_BYTES) -> str:
    """Bound persisted/returned text by bytes without leaving broken UTF-8."""
    encoded = value.encode("utf-8")
    if len(encoded) <= max_bytes:
        return value
    return encoded[:max_bytes].decode("utf-8", errors="ignore")


def _require_global_address(value: str) -> None:
    try:
        address = ipaddress.ip_address(value.split("%", 1)[0])
    except ValueError as error:
        raise SourceIngestionError("URL resolved to an invalid address") from error
    if not address.is_global:
        raise SourceIngestionError("URL must resolve only to public addresses")


async def _default_resolver(hostname: str, port: int) -> list[str]:
    loop = asyncio.get_running_loop()
    try:
        records = await loop.getaddrinfo(
            hostname,
            port,
            type=socket.SOCK_STREAM,
        )
    except socket.gaierror as error:
        raise SourceIngestionError("URL hostname could not be resolved") from error
    return sorted({record[4][0] for record in records})


async def validate_public_url(
    url: str,
    *,
    resolver: Resolver = _default_resolver,
) -> str:
    """Validate syntax and require every resolved address to be globally routable."""
    return (await resolve_public_target(url, resolver=resolver)).url


async def resolve_public_target(
    url: str,
    *,
    resolver: Resolver = _default_resolver,
) -> ValidatedTarget:
    """Validate a URL and bind it to one approved address for the next request."""
    if not isinstance(url, str) or len(url) > 2_048:
        raise SourceIngestionError("URL is invalid or too long")
    if url != url.strip() or any(ord(character) < 32 for character in url):
        raise SourceIngestionError("URL contains invalid whitespace")
    try:
        parsed = urlsplit(url)
        port = parsed.port
    except ValueError as error:
        raise SourceIngestionError("URL contains an invalid port") from error
    if parsed.scheme.lower() not in {"http", "https"}:
        raise SourceIngestionError("Only http and https URLs are supported")
    if not parsed.hostname or parsed.username is not None or parsed.password is not None:
        raise SourceIngestionError("URL credentials are not allowed")
    if parsed.fragment:
        # Fragments never reach the server and are removed from the canonical target.
        parsed = parsed._replace(fragment="")
        url = parsed.geturl()
    port = (443 if parsed.scheme.lower() == "https" else 80) if port is None else port
    if port < 1 or port > 65_535:
        raise SourceIngestionError("URL contains an invalid port")

    try:
        literal = ipaddress.ip_address(parsed.hostname.split("%", 1)[0])
    except ValueError:
        addresses = await resolver(parsed.hostname, port)
        if not addresses:
            raise SourceIngestionError("URL hostname did not resolve")
        for address in addresses:
            _require_global_address(address)
        pinned_ip = addresses[0]
    else:
        _require_global_address(str(literal))
        pinned_ip = str(literal)
    hostname = parsed.hostname.encode("idna").decode("ascii")
    return ValidatedTarget(url=url, hostname=hostname, pinned_ip=pinned_ip)


async def fetch_public_text(
    url: str,
    *,
    resolver: Resolver = _default_resolver,
    transport: Optional[httpx.AsyncBaseTransport] = None,
    max_bytes: int = MAX_FETCH_BYTES,
    max_redirects: int = MAX_REDIRECTS,
) -> FetchedText:
    """Fetch safe textual content while validating every redirect target."""
    current_url = url
    timeout = httpx.Timeout(connect=5.0, read=10.0, write=5.0, pool=5.0)
    for redirect_count in range(max_redirects + 1):
        target = await resolve_public_target(current_url, resolver=resolver)
        current_url = target.url
        hop_transport = transport or PinnedAsyncHTTPTransport(
            hostname=target.hostname,
            pinned_ip=target.pinned_ip,
        )
        async with httpx.AsyncClient(
            timeout=timeout,
            follow_redirects=False,
            transport=hop_transport,
            trust_env=False,
        ) as client:
            try:
                async with client.stream(
                    "GET",
                    current_url,
                    headers={"User-Agent": "Plotline/1.0 source reader"},
                ) as response:
                    if response.status_code in REDIRECT_STATUSES:
                        location = response.headers.get("location")
                        if not location:
                            raise SourceIngestionError("Redirect response omitted its target")
                        if redirect_count >= max_redirects:
                            raise SourceIngestionError("URL redirected too many times")
                        current_url = urljoin(current_url, location)
                        continue
                    if response.status_code < 200 or response.status_code >= 300:
                        raise SourceIngestionError(
                            f"URL returned HTTP {response.status_code}"
                        )
                    content_type = response.headers.get("content-type", "")
                    media_type = content_type.split(";", 1)[0].strip().lower()
                    if media_type not in SAFE_CONTENT_TYPES:
                        raise SourceIngestionError("URL content type is not supported")
                    declared_length = response.headers.get("content-length")
                    if declared_length and declared_length.isdigit() and int(declared_length) > max_bytes:
                        raise SourceIngestionError("URL content is too large")
                    body = bytearray()
                    async for chunk in response.aiter_bytes():
                        if len(body) + len(chunk) > max_bytes:
                            raise SourceIngestionError("URL content is too large")
                        body.extend(chunk)
                    encoding = response.encoding or "utf-8"
                    try:
                        text = bytes(body).decode(encoding, errors="replace")
                    except LookupError:
                        text = bytes(body).decode("utf-8", errors="replace")
                    return FetchedText(
                        final_url=current_url,
                        content_type=media_type,
                        text=text,
                    )
            except httpx.RequestError as error:
                raise SourceIngestionError("URL request failed") from error
    raise SourceIngestionError("URL redirected too many times")


def ensure_contained(path: Path, root: Path) -> Path:
    resolved_root = root.resolve()
    resolved = path.resolve()
    if not resolved.is_relative_to(resolved_root):
        raise SourceIngestionError("Source path escaped the project directory")
    return resolved


def validate_docx_archive(
    path: Path,
    *,
    max_files: int = MAX_DOCX_FILES,
    max_uncompressed_bytes: int = MAX_DOCX_UNCOMPRESSED_BYTES,
) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise SourceIngestionError("DOCX signature is invalid")
            if len(entries) > max_files:
                raise SourceIngestionError("DOCX contains too many files")
            expanded = sum(entry.file_size for entry in entries)
            if expanded > max_uncompressed_bytes:
                raise SourceIngestionError("DOCX expanded size is too large")
            for entry in entries:
                if entry.file_size > max_uncompressed_bytes:
                    raise SourceIngestionError("DOCX expanded size is too large")
                if entry.compress_size == 0 and entry.file_size > 0:
                    raise SourceIngestionError("DOCX compression ratio is unsafe")
                if entry.compress_size and entry.file_size / entry.compress_size > 200:
                    raise SourceIngestionError("DOCX compression ratio is unsafe")
    except zipfile.BadZipFile as error:
        raise SourceIngestionError("DOCX signature is invalid") from error


def validate_upload_signature(path: Path, extension: str) -> None:
    with path.open("rb") as handle:
        prefix = handle.read(8)
    if extension == ".pdf" and not prefix.startswith(b"%PDF-"):
        raise SourceIngestionError("PDF signature does not match its extension")
    if extension == ".docx":
        if not prefix.startswith(b"PK"):
            raise SourceIngestionError("DOCX signature does not match its extension")
        validate_docx_archive(path)
    if extension in {".txt", ".md"}:
        if b"\x00" in prefix:
            raise SourceIngestionError("Text source contains binary data")


def _apply_worker_limits(cpu_seconds: int, memory_bytes: int) -> None:
    try:
        import resource

        resource.setrlimit(resource.RLIMIT_CPU, (cpu_seconds, cpu_seconds + 1))
        resource.setrlimit(resource.RLIMIT_FSIZE, (MAX_UPLOAD_BYTES, MAX_UPLOAD_BYTES))
        # RLIMIT_AS is not reliably enforced by Darwin. Linux provides the
        # hard address-space boundary used in production containers.
        if sys.platform.startswith("linux"):
            resource.setrlimit(resource.RLIMIT_AS, (memory_bytes, memory_bytes))
    except (ImportError, OSError, ValueError):
        pass


def _extract_source_content(
    kind: str,
    payload: str,
    max_chars: int,
    max_pdf_pages: int,
    max_pdf_objects: int,
    worker_delay_seconds: float,
) -> str | tuple[str, str]:
    if worker_delay_seconds:
        time.sleep(worker_delay_seconds)
    if kind == "pdf":
        import PyPDF2

        with Path(payload).open("rb") as source:
            reader = PyPDF2.PdfReader(source)
            if len(reader.pages) > max_pdf_pages:
                raise SourceIngestionError("PDF contains too many pages")
            object_count = int(reader.trailer.get("/Size", 0) or 0)
            if object_count > max_pdf_objects:
                raise SourceIngestionError("PDF contains too many objects")
            extracted = "".join(
                page.extract_text() or "" for page in reader.pages
            ).strip()
    elif kind == "docx":
        from docx import Document

        path = Path(payload)
        validate_docx_archive(path)
        document = Document(path)
        if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
            raise SourceIngestionError("DOCX contains too many paragraphs")
        extracted = "\n".join(
            paragraph.text for paragraph in document.paragraphs
        ).strip()
    elif kind == "text":
        extracted = Path(payload).read_text(encoding="utf-8", errors="replace")
    elif kind == "html":
        from app.utils.file_extraction import extract_text_from_html

        title, text_content = extract_text_from_html(payload)
        return title[:512], truncate_utf8(text_content[:max_chars])
    else:
        raise SourceIngestionError("Unsupported extraction type")
    return truncate_utf8(extracted[:max_chars])


def _extraction_worker(
    output,
    kind: str,
    payload: str,
    max_chars: int,
    max_pdf_pages: int,
    max_pdf_objects: int,
    worker_delay_seconds: float,
) -> None:
    try:
        _apply_worker_limits(cpu_seconds=5, memory_bytes=512 * 1024 * 1024)
        extracted = _extract_source_content(
            kind,
            payload,
            max_chars,
            max_pdf_pages,
            max_pdf_objects,
            worker_delay_seconds,
        )
        output.send((True, extracted))
    except BaseException as error:
        try:
            output.send((False, str(error) or "Source extraction failed"))
        except (BrokenPipeError, EOFError, OSError):
            pass
    finally:
        output.close()


def _stop_extraction_process(process) -> None:
    if not process.is_alive():
        process.join()
        return
    process.terminate()
    process.join(0.1)
    if process.is_alive():
        process.kill()
        process.join(0.1)


def _run_extraction_process(
    kind: str,
    payload: str,
    *,
    timeout_seconds: float,
    max_chars: int,
    max_pdf_pages: int,
    max_pdf_objects: int,
    worker_delay_seconds: float,
):
    context = multiprocessing.get_context("spawn")
    receiver, sender = context.Pipe(duplex=False)
    process = context.Process(
        target=_extraction_worker,
        args=(
            sender,
            kind,
            payload,
            max_chars,
            max_pdf_pages,
            max_pdf_objects,
            worker_delay_seconds,
        ),
    )
    deadline = time.monotonic() + timeout_seconds
    try:
        process.start()
    except BaseException:
        receiver.close()
        sender.close()
        raise
    sender.close()

    received = queue.Queue(maxsize=1)

    def drain_result() -> None:
        try:
            received.put((True, receiver.recv()))
        except (EOFError, OSError) as error:
            received.put((False, error))

    reader = threading.Thread(target=drain_result, daemon=True)
    reader.start()
    try:
        try:
            remaining = max(0, deadline - time.monotonic())
            has_message, message = received.get(timeout=remaining)
        except queue.Empty as error:
            _stop_extraction_process(process)
            raise SourceIngestionError("Source extraction timed out") from error
        if not has_message:
            _stop_extraction_process(process)
            raise SourceIngestionError("Source extraction process failed") from message

        remaining = max(0, deadline - time.monotonic())
        process.join(remaining)
        if process.is_alive():
            _stop_extraction_process(process)
            raise SourceIngestionError("Source extraction timed out")

        succeeded, result = message
        if process.exitcode != 0 and succeeded:
            raise SourceIngestionError("Source extraction process failed")
        if not succeeded:
            raise SourceIngestionError(result)
        return result
    finally:
        receiver.close()
        reader.join(0.1)


async def extract_source_in_subprocess(
    kind: Literal["pdf", "docx", "text", "html"],
    source: Path | str,
    *,
    timeout_seconds: float = 10.0,
    max_chars: int = MAX_EXTRACTED_CHARS,
    max_pdf_pages: int = MAX_PDF_PAGES,
    max_pdf_objects: int = MAX_PDF_OBJECTS,
    worker_delay_seconds: float = 0,
):
    """Parse hostile formats in a process that can be forcefully terminated."""
    payload = str(source)
    if os.getenv("VERCEL"):
        try:
            return await asyncio.wait_for(
                asyncio.to_thread(
                    _extract_source_content,
                    kind,
                    payload,
                    max_chars,
                    max_pdf_pages,
                    max_pdf_objects,
                    worker_delay_seconds,
                ),
                timeout=timeout_seconds,
            )
        except asyncio.TimeoutError as error:
            raise SourceIngestionError("Source extraction timed out") from error
    return await asyncio.to_thread(
        _run_extraction_process,
        kind,
        payload,
        timeout_seconds=timeout_seconds,
        max_chars=max_chars,
        max_pdf_pages=max_pdf_pages,
        max_pdf_objects=max_pdf_objects,
        worker_delay_seconds=worker_delay_seconds,
    )
